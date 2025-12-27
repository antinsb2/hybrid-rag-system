"""
Document loaders for multiple file formats.
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
import PyPDF2
import docx
from bs4 import BeautifulSoup


@dataclass
class Document:
    """Represents a loaded document with metadata."""
    content: str
    source: str
    metadata: dict


class BaseLoader(ABC):
    """Base class for document loaders."""
    
    @abstractmethod
    def load(self, file_path: str) -> Document:
        """Load document from file."""
        pass


class PDFLoader(BaseLoader):
    """Load PDF documents."""
    
    def load(self, file_path: str) -> Document:
        """Extract text from PDF."""
        path = Path(file_path)
        
        with open(path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            # Extract text from all pages
            text = []
            for page in reader.pages:
                text.append(page.extract_text())
            
            content = '\n\n'.join(text)
            
            # Extract metadata
            metadata = {
                'source': str(path),
                'type': 'pdf',
                'pages': len(reader.pages),
            }
            
            if reader.metadata:
                metadata.update({
                    'author': reader.metadata.get('/Author', ''),
                    'title': reader.metadata.get('/Title', ''),
                    'created': reader.metadata.get('/CreationDate', ''),
                })
            
            return Document(
                content=content,
                source=str(path),
                metadata=metadata
            )


class DOCXLoader(BaseLoader):
    """Load DOCX documents."""
    
    def load(self, file_path: str) -> Document:
        """Extract text from DOCX."""
        path = Path(file_path)
        doc = docx.Document(path)
        
        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n\n'.join(paragraphs)
        
        metadata = {
            'source': str(path),
            'type': 'docx',
            'paragraphs': len(paragraphs),
        }
        
        return Document(
            content=content,
            source=str(path),
            metadata=metadata
        )


class HTMLLoader(BaseLoader):
    """Load HTML documents."""
    
    def load(self, file_path: str) -> Document:
        """Extract text from HTML."""
        path = Path(file_path)
        
        with open(path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
            
            # Remove script and style tags
            for tag in soup(['script', 'style']):
                tag.decompose()
            
            # Extract text
            content = soup.get_text(separator='\n', strip=True)
            
            metadata = {
                'source': str(path),
                'type': 'html',
                'title': soup.title.string if soup.title else '',
            }
            
            return Document(
                content=content,
                source=str(path),
                metadata=metadata
            )


class TextLoader(BaseLoader):
    """Load plain text and markdown documents."""
    
    def load(self, file_path: str) -> Document:
        """Load text file."""
        path = Path(file_path)
        
        with open(path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        metadata = {
            'source': str(path),
            'type': 'text',
            'extension': path.suffix,
        }
        
        return Document(
            content=content,
            source=str(path),
            metadata=metadata
        )


class DocumentLoader:
    """Main document loader that routes to appropriate loader."""
    
    def __init__(self):
        self.loaders = {
            '.pdf': PDFLoader(),
            '.docx': DOCXLoader(),
            '.html': HTMLLoader(),
            '.htm': HTMLLoader(),
            '.txt': TextLoader(),
            '.md': TextLoader(),
        }
    
    def load(self, file_path: str) -> Document:
        """
        Load document from file.
        
        Args:
            file_path: Path to document
            
        Returns:
            Document object with content and metadata
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension not in self.loaders:
            raise ValueError(f"Unsupported file type: {extension}")
        
        return self.loaders[extension].load(file_path)
